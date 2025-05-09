from flask import Flask, render_template, request, redirect, url_for, session, send_file, jsonify
import openai
import os
import logging
import sys
import threading
import subprocess
from logging.handlers import RotatingFileHandler
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
import secrets
from markupsafe import escape
from config import get_config
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
# Security imports
from flask_talisman import Talisman
import jwt
import bcrypt
import re
from forms import PromptForm
from urllib.parse import urlparse
from unittest.mock import MagicMock

# Initialize logging
def setup_logging():
    if not os.path.exists('logs'):
        os.mkdir('logs')
    file_handler = RotatingFileHandler('logs/solar_assistant.log', maxBytes=10240, backupCount=5)
    file_handler.setFormatter(logging.Formatter(
        '%(asctime)s %(levelname)s: %(message)s [in %(pathname)s:%(lineno)d]'
    ))
    file_handler.setLevel(logging.INFO)
    
    app_logger = logging.getLogger('solar_assistant')
    app_logger.setLevel(logging.INFO)
    app_logger.addHandler(file_handler)
    
    return app_logger

logger = setup_logging()

def create_app(config_name=None):
    # Load environment variables
    load_dotenv()

    # Create and configure the app
    app = Flask(__name__)
    
    # Load config from object
    app.config.from_object(get_config(config_name))
    
    # Important: Set the secret key early and explicitly
    secret_key = app.config.get('SECRET_KEY') or os.getenv('FLASK_SECRET') or secrets.token_hex(16)
    app.secret_key = secret_key
    
    # Special handling for test environment
    testing = config_name == "testing" or app.config.get('TESTING', False)
    if testing:
        app.config['TESTING'] = True
        app.config['WTF_CSRF_ENABLED'] = False
        # Ensure we have a secret key for testing
        if not app.secret_key or app.secret_key == 'dev-key-NOT-FOR-PRODUCTION-USE':
            app.secret_key = 'test-secret-key-for-pytest'
        
        # Skip Talisman and other security middleware in testing
        logger.info(f"Testing mode detected. Skipping security middleware.")
    else:
        # Set up Flask-Talisman for security headers (only in non-testing environments)
        csp = {
            'default-src': "'self'",
            'style-src': ["'self'", "'unsafe-inline'"],
            'script-src': ["'self'"],
            'img-src': ["'self'", "data:"],
            'font-src': ["'self'"],
            'connect-src': ["'self'", "mailto:", "https://www.linkedin.com", "https://www.facebook.com", 
                            "https://api.whatsapp.com", "https://docs.google.com"],
            'form-action': ["'self'", "mailto:", "https://www.linkedin.com", "https://www.facebook.com", 
                            "https://api.whatsapp.com", "https://docs.google.com"]
        }
        
        # Only enable HTTPS in production
        force_https = app.config.get('ENV', 'production') == 'production'
        if not testing:
            Talisman(app, content_security_policy=csp, force_https=force_https)
    
    # Manually add CORS headers instead of using flask_cors
    @app.after_request
    def add_cors_headers(response):
        response.headers['Access-Control-Allow-Origin'] = '*'
        response.headers['Access-Control-Allow-Methods'] = 'GET, POST, PUT, DELETE, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type, Authorization'
        return response
    
    # Set up rate limiting 
    limiter = Limiter(
        app=app,
        key_func=get_remote_address,
        default_limits=["200 per day", "50 per hour"],
        storage_uri="memory://",
    )
    
    # Health check endpoints
    @app.route('/health')
    @limiter.exempt  # Don't rate limit health checks
    def health_check():
        return jsonify({"status": "healthy", "version": "1.0.0"})
    
    @app.route("/healthz")
    @limiter.exempt  # Don't rate limit health checks
    def healthz():
        """Health check endpoint for Render monitoring."""
        return jsonify({"status": "healthy"}), 200

    # Initialize OpenAI client with better error handling for tests
    # Make client available at app level so routes can access it
    app.openai_client = None
    
    try:
        api_key = app.config.get('OPENAI_API_KEY') or os.getenv('OPENAI_API_KEY')
        if not api_key:
            if not app.config.get('TESTING'):
                logger.error("OpenAI API key is missing")
                raise ValueError("OpenAI API key is required")
            else:
                # For testing, use a dummy key
                api_key = "test-key"
                logger.info("Using test API key in testing mode")
        
        # Try both client styles based on what's available
        # New style client (OpenAI v1.0.0+)
        if hasattr(openai, "OpenAI"):
            try:
                app.openai_client = openai.OpenAI(api_key=api_key)
                logger.info("OpenAI client initialized via class")
            except Exception as e:
                if not app.config.get('TESTING'):
                    logger.error(f"Error initializing OpenAI client: {e}")
                    raise
        
        # Legacy style fallback
        if app.openai_client is None:
            try:
                openai.api_key = api_key
                app.openai_client = openai
                logger.info("OpenAI client set to module-level API")
            except Exception as e:
                if not app.config.get('TESTING'):
                    logger.error(f"Error initializing legacy OpenAI client: {e}")
                    raise
        
        # For testing, create a basic client stub if needed
        if app.config.get('TESTING') and app.openai_client is None:
            app.openai_client = MagicMock()
            logger.info("Created mock OpenAI client for testing")
            
    except Exception as e:
        if not app.config.get('TESTING'):
            logger.error(f"Error initializing OpenAI client: {e}")
            import traceback
            logger.error(traceback.format_exc())
        else:
            # For testing, create a basic client stub
            app.openai_client = MagicMock()
            logger.info("Created mock OpenAI client for testing after error")

    # Add security middleware
    @app.before_request
    def security_checks():
        # Validate request data to prevent injection attacks
        if request.method == 'POST':
            # Check for suspicious patterns in form data
            for key, value in request.form.items():
                if isinstance(value, str) and contains_suspicious_patterns(value):
                    logger.warning(f"Potentially malicious input detected from {request.remote_addr}")
                    return render_template('error.html', error="Invalid input detected"), 400
        
        # Prevent clickjacking by checking Referer
        if request.method != 'GET' and not is_safe_referrer(request.referrer):
            logger.warning(f"Potential CSRF attempt from {request.remote_addr} with referrer {request.referrer}")
            return render_template('error.html', error="Invalid request origin"), 403
    
    # Error handlers
    @app.errorhandler(404)
    def not_found_error(error):
        return render_template('error.html', error="Page not found"), 404
    
    @app.errorhandler(500)
    def internal_error(error):
        logger.error(f"Server Error: {error}")
        return render_template('error.html', error="An internal error occurred"), 500
    
    @app.errorhandler(429)
    def ratelimit_handler(e):
        logger.warning(f"Rate limit exceeded: {e.description}")
        return render_template('error.html', error=f"Rate limit exceeded. {e.description}"), 429
    
    @app.errorhandler(Exception)
    def unhandled_exception(e):
        logger.error(f"Unhandled Exception: {e}")
        if app.config['DEBUG']:
            # In debug mode, let the default handler deal with it
            raise e
        return render_template('error.html', error="An unexpected error occurred"), 500
    
    # Routes
    @app.route('/', methods=['GET', 'POST'])
    def index():
        form = PromptForm()
        response_text = ""
        error_message = None

        # If client not set, show error
        if app.openai_client is None:
            error_message = "OpenAI service is currently unavailable. Please try again later."
            logger.error("Route accessed with uninitialized OpenAI client")
            return render_template('index.html', form=form, response=session.get('response_text', ''), error=error_message)

        # Load the professional system prompt from the external file
        try:
            with open(app.config['PROMPT_PATH'], "r", encoding="utf-8") as f:
                system_prompt = f.read()
                logger.info(f"Successfully loaded prompt from {app.config['PROMPT_PATH']}")
        except FileNotFoundError:
            system_prompt = "You are a helpful solar PV system design assistant."
            error_message = "System configuration warning: Using default prompt."
        
        # Append structured report generation instructions
        report_instructions = """
1. Title and Definition
   • Start with a clear heading, e.g. “Charge Controller (Solar Charge Regulator)”
   • Provide a one-sentence definition of its role in a PV system.

2. Voltage Regulation
   • Explain how the controller prevents overcharge by regulating panel‑to‑battery voltage.
   • Describe reverse‑current protection at night.

3. Battery Health & Lifespan
   • Detail how correct charging preserves battery health and extends service life.
   • Cite common failure modes avoided by regulation.

4. Types of Controllers
   • PWM Controllers
     – Define Pulse Width Modulation
     – List ideal use‑cases (small 12V/24V systems, cost constraints)
   • MPPT Controllers
     – Define Maximum Power Point Tracking
     – Explain efficiency gains and voltage‑to‑amperage conversion
     – Highlight performance in variable light/temperature

5. Sizing & Compatibility
   • Show how to choose a controller based on system voltage (12V/24V/48V) and array short‑circuit current
   • Include a brief sizing calculation example.

6. Additional Features
   • Mention integrated LCD or LED status displays
   • List common safety protections (over‑temperature, reverse polarity)
   • Note networking and remote‑monitoring capabilities.

7. Summary & Recommendations
   • Recap why selecting the right controller is critical
   • Offer best‑practice tips (e.g. derating margin, vendor reliability)
"""
        system_prompt += "\n" + report_instructions

        # Append ProReport AI template for report formatting
        try:
            with open(app.config['TEMPLATE_PATH'], "r", encoding="utf-8") as tf:
                template_content = tf.read()
                logger.info(f"Loaded ProReport template from {app.config['TEMPLATE_PATH']}")
                system_prompt += "\n" + template_content
        except FileNotFoundError:
            logger.error(f"Template file not found: {app.config['TEMPLATE_PATH']}")

        if form.validate_on_submit():
            prompt = form.prompt.data
            # Adjust system prompt to respond in the selected language
            lang_map = {
                'en': 'English',
                'sw': 'Kiswahili',
                'ar': 'Arabic',
                'am': 'Amharic',
                'es': 'Spanish',
                'fr': 'French'
            }
            selected = form.language.data
            lang_name = lang_map.get(selected, 'English')
            # Instruct the AI to respond in the chosen language
            system_prompt = f"{system_prompt}\nPlease respond in {lang_name}."

            # Wrap actual chat call in try/except
            logger.info(f"Processing prompt of length {len(prompt)}")
            try:
                # Handle both new and legacy OpenAI client styles
                if hasattr(app.openai_client, 'chat') and hasattr(app.openai_client.chat, 'completions'):
                    # New style client (v1.0.0+)
                    response = app.openai_client.chat.completions.create(
                        model=app.config['OPENAI_MODEL'],
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": prompt}
                        ]
                    )
                    response_text = response.choices[0].message.content
                else:
                    # Legacy style client
                    response = app.openai_client.ChatCompletion.create(
                        model=app.config['OPENAI_MODEL'],
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": prompt}
                        ]
                    )
                    response_text = response.choices[0].message.content
                
                formatted_response = response_text.replace('\n', '<br>')
                session['response_text'] = response_text
                session['formatted_response'] = formatted_response
                logger.info(f"Successfully generated response of length {len(response_text)}")
            except Exception as e:
                logger.error(f"OpenAI API Error: {e}")
                error_message = f"Error generating response: {e}"
                response_text = "Sorry, there was an error processing your request."
                session['response_text'] = response_text
                session['formatted_response'] = response_text

        # Use the formatted response for display if available
        display_response = session.get('formatted_response', session.get('response_text', ''))

        return render_template('index.html', 
                              form=form,
                              response=display_response, 
                              error=error_message)
    
    @app.route('/robots.txt')
    def robots():
        """Serve robots.txt from static folder"""
        return app.send_static_file('robots.txt')
        
    @app.route('/favicon.ico')
    def favicon():
        """Serve favicon.ico from static folder"""
        return app.send_static_file('favicon.ico')
        
    @app.route('/clear')
    def clear():
        session.pop('response_text', None)
        session.pop('formatted_response', None)
        return redirect(url_for('index'))
    
    @app.route('/view-report', methods=['GET', 'POST'])
    def view_report():
        # Allow editing and saving updated report text
        if request.method == 'POST':
            updated = request.form.get('report_text', '')
            session['response_text'] = updated
        raw_text = session.get('response_text', 'No report available.')
        # Construct shareable URL for social/sharing
        share_url = request.url
        return render_template('view_report.html', response=raw_text, share_url=share_url)
    
    @app.route('/download-report')
    def download_report():
        response_text = session.get('response_text', 'No report available.')
        
        if response_text == 'No report available.':
            logger.warning("Attempted to download an empty report")
            return redirect(url_for('index'))
            
        try:
            doc = Document()
            doc.add_heading("Solar System Design Report", 0)
            
            for line in response_text.split('\n'):
                if line.strip().startswith("**") and line.strip().endswith("**"):
                    doc.add_heading(line.strip().strip("*"), level=1)
                else:
                    doc.add_paragraph(line.strip())
            
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            
            logger.info("Report document generated successfully")
            
            return send_file(
                file_stream,
                as_attachment=True,
                download_name="Solar_Report.docx",
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        except Exception as e:
            logger.error(f"Error generating report document: {e}")
            return render_template('error.html', error="Error generating document"), 500
    
    return app

def contains_suspicious_patterns(value):
    """Check for potentially malicious patterns in input"""
    # Check for common SQL injection or XSS patterns
    suspicious_patterns = [
        r"(?i)(<script|javascript:|on\w+=|\balert\(|\beval\()",  # Basic XSS
        r"(?i)(union\s+select|insert\s+into|update\s+.*?\s+set|delete\s+from)",  # SQL injection
        r"(?i)(\.\.\/|\.\.\\|\/etc\/passwd)"  # Path traversal
    ]
    
    return any(re.search(pattern, value) for pattern in suspicious_patterns)

def is_safe_referrer(referrer):
    """Validate that request comes from an allowed origin"""
    if not referrer:
        return True

    # Start with localhost for dev
    allowed_hosts = ['localhost', '127.0.0.1']

    # Add any hosts you want via comma‑separated env var
    app_hosts = os.environ.get('ALLOWED_APP_HOSTS', '')
    if app_hosts:
        allowed_hosts += [h.strip() for h in app_hosts.split(',')]

    # Auto‑include the Render domain when deployed
    render_host = os.environ.get('RENDER_EXTERNAL_HOSTNAME')
    if render_host:
        allowed_hosts.append(render_host)
        # Also allow www.<host>
        if not render_host.startswith('www.'):
            allowed_hosts.append(f'www.{render_host}')

    # Finally, check the referrer against our list
    logger.debug(f"Allowed hosts: {allowed_hosts}, Referrer: {referrer}")
    ref_host = urlparse(referrer).hostname
    return ref_host in allowed_hosts

def check_for_updates():
    """Run the dependency update checker in a subprocess"""
    try:
        # Start the update process dialog
        print("\n" + "-" * 60)
        print("Checking for dependency updates...")
        print("-" * 60)
        
        # Execute the update script and pass its output to the terminal
        update_script = os.path.join(os.path.dirname(os.path.abspath(__file__)), "update_dependencies.py")
        subprocess.run([sys.executable, update_script], check=False)
        
        print("-" * 60 + "\n")
    except Exception as e:
        print(f"Error checking for updates: {e}")

# Application entry point
if __name__ == '__main__':
    # Check for updates before starting the app
    update_thread = threading.Thread(target=check_for_updates)
    update_thread.daemon = True
    update_thread.start()
    
    app = create_app()
    # The following line has been removed for Render deployment
    # port = int(os.getenv('PORT', 8003))
    # app.run(debug=app.config['DEBUG'], port=port)
else:
    # Create app instance for gunicorn to import
    app = create_app('production')

