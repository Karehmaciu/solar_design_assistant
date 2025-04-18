from flask import Flask, render_template, request, redirect, url_for, session, send_file, jsonify
import os
import logging
import sys
import threading
import subprocess
from io import BytesIO
from docx import Document
from dotenv import load_dotenv
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_talisman import Talisman
from flask_cors import CORS
import re
from forms import PromptForm
from solar_design_agent import SolarDesignAgent, setup_logging

# Initialize logging
logger = setup_logging()

def create_app(config_name=None):
    # Load environment variables
    load_dotenv()
    
    # Create and configure the app
    app = Flask(__name__)
    
    # Configure secret key from environment or generate a random one
    app.config['SECRET_KEY'] = os.getenv('FLASK_SECRET', os.urandom(24).hex())
    app.config['DEBUG'] = os.getenv('FLASK_ENV', 'production') != 'production'
    app.config['OPENAI_API_KEY'] = os.getenv('OPENAI_API_KEY')
    app.config['OPENAI_MODEL'] = os.getenv('OPENAI_MODEL', 'gpt-4')
    app.config['PROMPT_PATH'] = os.getenv('PROMPT_PATH', 'prompts/kbs_solar_prompt_final.txt')
    app.config['TEMPLATE_PATH'] = os.getenv('TEMPLATE_PATH', 'prompts/proreport_template.md')
    
    # Security enhancements
    # Set up Flask-Talisman for security headers (CSP, HSTS, etc.)
    csp = {
        'default-src': "'self'",
        'style-src': ["'self'", "'unsafe-inline'"],  # Allow inline styles for simplicity
        'script-src': ["'self'"],
        'img-src': ["'self'", "data:"],
        'font-src': ["'self'"],
        'connect-src': ["'self'", "mailto:", "https://www.linkedin.com", "https://www.facebook.com", 
                        "https://api.whatsapp.com", "https://docs.google.com"],
        'form-action': ["'self'", "mailto:", "https://www.linkedin.com", "https://www.facebook.com", 
                        "https://api.whatsapp.com", "https://docs.google.com"]
    }
    
    # Only enable HTTPS in production
    force_https = app.config.get('ENV', 'production') == 'production'
    Talisman(app, content_security_policy=csp, force_https=force_https)
    
    # Configure CORS to restrict access to trusted domains
    CORS(app, resources={r"/*": {"origins": ["http://localhost:*", "https://*.solar-assistant.example.com"]}})
    
    # Set up rate limiting
    limiter = Limiter(
        app=app,
        key_func=get_remote_address,
        default_limits=["200 per day", "50 per hour"],
        storage_uri="memory://",
    )
    
    # Initialize the Solar Design Agent
    try:
        api_key = app.config.get('OPENAI_API_KEY')
        solar_agent = SolarDesignAgent(api_key, model=app.config.get('OPENAI_MODEL'))
        logger.info("Solar Design Agent initialized successfully")
    except Exception as e:
        logger.error(f"Error initializing Solar Design Agent: {e}")
        import traceback; logger.error(traceback.format_exc())
        solar_agent = None  # Handle None in routes
    
    # Add security middleware
    @app.before_request
    def security_checks():
        # Validate request data to prevent injection attacks
        if request.method == 'POST':
            # Check for suspicious patterns in form data
            for key, value in request.form.items():
                if isinstance(value, str) and contains_suspicious_patterns(value):
                    logger.warning(f"Potentially malicious input detected from {request.remote_addr}")
                    return render_template('error.html', error="Invalid input detected"), 400
        
        # Prevent clickjacking by checking Referer
        if request.method != 'GET' and not is_safe_referrer(request.referrer):
            logger.warning(f"Potential CSRF attempt from {request.remote_addr} with referrer {request.referrer}")
            return render_template('error.html', error="Invalid request origin"), 403
    
    # Add health check endpoints
    @app.route('/health')
    @limiter.exempt  # Don't rate limit health checks
    def health_check():
        return jsonify({"status": "healthy", "version": "1.0.0"})
    
    @app.route("/healthz")
    @limiter.exempt  # Don't rate limit health checks
    def healthz():
        """Health check endpoint for Render monitoring."""
        return jsonify({"status": "healthy"}), 200

    # Error handlers
    @app.errorhandler(404)
    def not_found_error(error):
        return render_template('error.html', error="Page not found"), 404
    
    @app.errorhandler(500)
    def internal_error(error):
        logger.error(f"Server Error: {error}")
        return render_template('error.html', error="An internal error occurred"), 500
    
    @app.errorhandler(429)
    def ratelimit_handler(e):
        logger.warning(f"Rate limit exceeded: {e.description}")
        return render_template('error.html', error=f"Rate limit exceeded. {e.description}"), 429
    
    @app.errorhandler(Exception)
    def unhandled_exception(e):
        logger.error(f"Unhandled Exception: {e}")
        if app.config['DEBUG']:
            # In debug mode, let the default handler deal with it
            raise e
        return render_template('error.html', error="An unexpected error occurred"), 500
    
    # Routes
    @app.route('/', methods=['GET', 'POST'])
    def index():
        form = PromptForm()
        response_text = ""
        error_message = None
        
        # Check if Solar Design Agent is properly initialized
        if solar_agent is None:
            error_message = "OpenAI service is currently unavailable. Please try again later."
            logger.error("Route accessed with uninitialized Solar Design Agent")
            return render_template('index.html', 
                                 form=form,
                                 response=session.get('response_text', ''), 
                                 error=error_message)
        
        # Load system prompt and template
        system_prompt = solar_agent.load_system_prompt(app.config['PROMPT_PATH'])
        template = solar_agent.load_template(app.config['TEMPLATE_PATH'])
        
        if form.validate_on_submit():
            # Input validation is handled by WTForms
            prompt = form.prompt.data
            
            # Generate response using the Solar Design Agent
            response_text = solar_agent.generate_response(prompt, system_prompt, template)
            
            # Convert Markdown line breaks to HTML <br> tags for proper display
            # This preserves formatting in the HTML response
            formatted_response = response_text.replace('\n', '<br>')
            
            # Store both versions - original for download, formatted for display
            session['response_text'] = response_text
            session['formatted_response'] = formatted_response
        
        # Use the formatted response for display if available
        display_response = session.get('formatted_response', session.get('response_text', ''))
        
        return render_template('index.html', 
                              form=form,
                              response=display_response, 
                              error=error_message)
    
    @app.route('/clear')
    def clear():
        session.pop('response_text', None)
        session.pop('formatted_response', None)
        return redirect(url_for('index'))
    
    @app.route('/view-report', methods=['GET', 'POST'])
    def view_report():
        # Allow editing and saving updated report text
        if request.method == 'POST':
            updated = request.form.get('report_text', '')
            session['response_text'] = updated
        raw_text = session.get('response_text', 'No report available.')
        # Construct shareable URL for social/sharing
        share_url = request.url
        return render_template('view_report.html', response=raw_text, share_url=share_url)
    
    @app.route('/download-report')
    def download_report():
        response_text = session.get('response_text', 'No report available.')
        
        if response_text == 'No report available.':
            logger.warning("Attempted to download an empty report")
            return redirect(url_for('index'))
            
        try:
            doc = Document()
            doc.add_heading("Solar System Design Report", 0)
            
            for line in response_text.split('\n'):
                if line.strip().startswith("**") and line.strip().endswith("**"):
                    doc.add_heading(line.strip().strip("*"), level=1)
                else:
                    doc.add_paragraph(line.strip())
            
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            
            logger.info("Report document generated successfully")
            
            return send_file(
                file_stream,
                as_attachment=True,
                download_name="Solar_Report.docx",
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        except Exception as e:
            logger.error(f"Error generating report document: {e}")
            return render_template('error.html', error="Error generating document"), 500
    
    return app

def contains_suspicious_patterns(value):
    """Check for potentially malicious patterns in input"""
    # Check for common SQL injection or XSS patterns
    suspicious_patterns = [
        r"(?i)(<script|javascript:|on\w+=|\balert\(|\beval\()",  # Basic XSS
        r"(?i)(union\s+select|insert\s+into|update\s+.*?\s+set|delete\s+from)",  # SQL injection
        r"(?i)(\.\.\/|\.\.\\|\/etc\/passwd)"  # Path traversal
    ]
    
    return any(re.search(pattern, value) for pattern in suspicious_patterns)

def is_safe_referrer(referrer):
    """Validate that request comes from an allowed origin"""
    if not referrer:
        return True  # No referrer might be OK depending on your security posture
    
    allowed_hosts = ['localhost', '127.0.0.1', 'solar-assistant.example.com']
    for host in allowed_hosts:
        if host in referrer:
            return True
    
    return False

def check_for_updates():
    """Run the dependency update checker in a subprocess"""
    try:
        # Start the update process dialog
        print("\n" + "-" * 60)
        print("Checking for dependency updates...")
        print("-" * 60)
        
        # Execute the update script and pass its output to the terminal
        update_script = os.path.join(os.path.dirname(os.path.abspath(__file__)), "update_dependencies.py")
        subprocess.run([sys.executable, update_script], check=False)
        
        print("-" * 60 + "\n")
    except Exception as e:
        print(f"Error checking for updates: {e}")

# Application entry point
if __name__ == '__main__':
    # Check for updates before starting the app
    update_thread = threading.Thread(target=check_for_updates)
    update_thread.daemon = True
    update_thread.start()
    
    app = create_app()
    port = int(os.getenv('PORT', 8003))
    app.run(debug=app.config['DEBUG'], port=port)